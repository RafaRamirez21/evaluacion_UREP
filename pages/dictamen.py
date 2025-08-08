import pandas as pd
import streamlit as st
from PIL import Image
import numpy as np
import io
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuración de la página
st.set_page_config(page_title="Evaluación UREP", layout="wide")

if "is_logged_in" not in st.session_state or not st.session_state["is_logged_in"]:
    st.switch_page("principal.py")

# Cargar logo
logo = Image.open("logo_2.png")
st.sidebar.image(logo, width=100, use_container_width=True)

# Cargar datos del Excel
excel_path = "requisitos_urep.xlsx"
df_documentos = pd.read_excel(excel_path, sheet_name="Documentos")
df_requisitos = pd.read_excel(excel_path, sheet_name="Requisitos")

# Filtrar documentos según el tipo de establecimiento seleccionado
tipo_establecimiento = st.session_state.get("seleccion", "")
documentos_filtrados = df_documentos[df_documentos["Tipos de establecimientos"] == tipo_establecimiento]
botones_sidebar = documentos_filtrados["Requisitos"].dropna().unique().tolist()


# --- CSS para botones ---
st.sidebar.markdown("""
<style>
    /* Estilos base para todos los botones */
    [data-testid="stButton"] button {
        border-radius: 8px !important;
        border: none !important;
        cursor: pointer !important;
        margin-bottom: 8px !important;
        width: 100% !important;
        text-align: center !important;
        font-size: 16px !important;
    }
    
    /* Botones de documentos (primarios) en el sidebar */
    [data-testid="stSidebar"] [data-testid="stButton"] button[kind="primary"] {
        background-color: #0056b3 !important;
        color: #fff !important;
        padding: 10px 18px !important;
        text-align: left !important;
    }
    
    /* Botones secundarios (Regresar y Cerrar sesión) */
    [data-testid="stSidebar"] [data-testid="stButton"] button[kind="secondary"] {
        background-color: #27F5E0 !important; /* Color turquesa */
        color: black !important;
        padding: 6px 10px !important;
    }
    
    /* Botón Generar Reportes (secondary) */
    [data-testid="stSidebar"] [data-testid="stButton"] button[data-testid="baseButton-generar"] {
        background-color: #FF6B6B !important; /* Color rojo claro */
        color: white !important;
        font-weight: bold !important;
    }
    
    /* Efecto hover para todos los botones */
    [data-testid="stButton"] button:hover {
        opacity: 0.9 !important;
    }
    
    /* Centrar los contenedores de los botones de descarga */
    .stDownloadButton {
        text-align: center !important;
    }
    
    /* Estilo específico para botones de descarga */
    .stDownloadButton button[data-testid="stBaseButton-secondary"] {
        background-color: #F0394D !important; /* Color turquesa */
        color: white !important;
        border-radius: 8px !important;
        border: none !important;
        padding: 6px 10px !important;
        font-size: 16px !important;
        cursor: pointer !important;
        margin-top: 5px !important;
        width: auto !important;
        display: block !important;
        margin-left: auto !important;
        margin-right: auto !important;
    }
</style>
""", unsafe_allow_html=True)

# Botones de navegación en el sidebar
space1, col1, col2, space = st.sidebar.columns([4,2,4,3],vertical_alignment="center")
with space1:
    if st.button("Cerrar sesión", key="cerrar", type="secondary", help="Cerrar sesión", use_container_width=True):
        st.session_state.clear()
        st.switch_page("principal.py")
with col2:
    if st.button("🢀", key="regresar", type="secondary", help="Regresar", use_container_width=True):
        st.switch_page("principal.py")
# En el archivo paste.txt, reemplaza la sección del botón "Nueva evaluación" con esto:

with space:
    if st.button("➕", key="nueva_evaluacion", type="secondary", help="Nueva evaluación", use_container_width=True):
        # Guardar las credenciales del usuario y estado de inicio de sesión
        claves_a_conservar = {
            "is_logged_in": st.session_state.get("is_logged_in", False),
            "password_correct": st.session_state.get("password_correct", False), 
            "user_id": st.session_state.get("user_id", ""),
            "user_name": st.session_state.get("user_name", ""),
            "user_role": st.session_state.get("user_role", ""),
            "auth_token": st.session_state.get("auth_token", "")
        }
        
        # Limpiar todos los estados
        st.session_state.clear()
        
        # Restaurar solo las credenciales del usuario
        for key, value in claves_a_conservar.items():
            if value:  # Solo restaurar si el valor no está vacío
                st.session_state[key] = value
        
        # Navegar a la página principal manteniendo la sesión
        st.switch_page("principal.py")




st.sidebar.markdown("---")
st.sidebar.title("Documentos a evaluar:")

# Persistencia de selección de documento y observaciones
if "documento_seleccionado" not in st.session_state:
    st.session_state["documento_seleccionado"] = None
if "observaciones_dict" not in st.session_state:
    st.session_state["observaciones_dict"] = {}
if "evaluaciones_documentos" not in st.session_state:
    st.session_state["evaluaciones_documentos"] = {}

# Definir documento_actual ANTES del bucle de botones
documento_actual = st.session_state.get("documento_seleccionado", None)

# Mostrar solo los botones funcionales
for doc in botones_sidebar:
    if st.sidebar.button(doc, key=f"btn_{doc}", type="primary"):
        # Guardar la evaluación actual antes de cambiar de documento
        if documento_actual:
            st.session_state["evaluaciones_documentos"][documento_actual] = st.session_state["observaciones_dict"].copy()
        # Cambiar al nuevo documento
        st.session_state["documento_seleccionado"] = doc
        # Cargar la evaluación previa si existe
        st.session_state["observaciones_dict"] = st.session_state["evaluaciones_documentos"].get(doc, {})
        # Actualizar documento_actual después de cambiar
        documento_actual = doc

# Mostrar usuario y datos principales
col_tramite,col_tipo, col_est, col_sol, col_usr = st.columns((3,3,3,1.5,1.5), gap="small", border=True)
with col_tramite:
    st.markdown(f"<span style='font-size: 14px; color:#005662'><strong><em>Tipo:</em></strong> {st.session_state.get('tipo_tramite', '')}</span>", unsafe_allow_html=True)
with col_tipo:
    st.markdown(f"<span style='font-size: 14px; color:#005662'><strong><em>Tipo:</em></strong> {st.session_state.get('seleccion', '')}</span>", unsafe_allow_html=True)
with col_est:
    st.markdown(f"<span style='font-size: 14px; color:#005662'><strong><em>Establecimiento:</em></strong> {st.session_state.get('nombre_establecimiento', '')}</span>", unsafe_allow_html=True)
with col_sol:
    st.markdown(f"<span style='font-size: 14px; color:#005662'><strong><em>Solicitud:</em></strong> {st.session_state.get('numero_solicitud', '')}</span>", unsafe_allow_html=True)
with col_usr:
    st.markdown(f"<span style='font-size: 14px; color:#005662'><strong><em>Usuario:</em></strong> {st.session_state.get('user_id', '')}</span>", unsafe_allow_html=True)
st.write("")

st.markdown(f"""
<div style="text-align: center; font-size: 28px; color:#005662">
    <strong><em>Evaluación UREP</em></strong>
</div>
""", unsafe_allow_html=True)
st.write("")

#st.write("Placas activas registradas:", st.session_state["placas_vehiculos"])

# Mostrar requisitos del documento seleccionado
if documento_actual:
    st.markdown(f"""
    <div style="text-align: center; font-size: 22px;">
        <strong>{documento_actual}</strong>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("---")

    requisitos_filtrados = df_requisitos[df_requisitos["Documento"] == documento_actual]

    for idx, row in requisitos_filtrados.iterrows():
        requisito = row["Requisito"]
        sugerencia = row["Observación precargada"] if "Observación precargada" in row else ""
        if isinstance(sugerencia, float) and np.isnan(sugerencia):
            sugerencia = ""

        radio_key = f"cumple_{documento_actual}_{idx}"
        obs_key = f"observacion_{documento_actual}_{idx}"

        cumple_valor = st.session_state["observaciones_dict"].get((documento_actual, requisito), {}).get("cumple", None)
        observacion_valor = st.session_state["observaciones_dict"].get((documento_actual, requisito), {}).get("observacion", "")

        st.markdown(f"""
        <div style="text-align: left; font-size: 18px;">
            <strong>{requisito}</strong>
        </div>
        """, unsafe_allow_html=True)

        cols = st.columns([1, 2])
        with cols[0]:
            cumple = st.radio(
                "",
                ("Cumple", "No Cumple"),
                key=radio_key,
                horizontal=True,
                index=None if cumple_valor is None else (0 if cumple_valor == "Cumple" else 1)
            )
        with cols[1]:
            if cumple == "No Cumple":
                # Si ya hay observación guardada y es distinta de la sugerida, mostrar esa, si no, sugerencia
                if observacion_valor and observacion_valor != sugerencia:
                    default_obs = observacion_valor
                else:
                    default_obs = sugerencia
                observacion = st.text_area(
                    "Observaciones",
                    value=default_obs,
                    key=obs_key
                )
                # Guardar la observación ingresada, ignorando la sugerida si el usuario la cambia
                st.session_state["observaciones_dict"][(documento_actual, requisito)] = {
                    "cumple": cumple,
                    "observacion": observacion
                }
            else:
                st.session_state["observaciones_dict"][(documento_actual, requisito)] = {
                    "cumple": cumple,
                    "observacion": ""
                }

def limpiar_requisito(requisito):
    import re
    return re.sub(r'^[^.]*\.\s*', '', requisito)

def int_to_roman(num):
    """Convierte un número entero a su representación en números romanos"""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num

def generar_observaciones_texto(df, documento_actual, contador_inicial=1):
    observaciones_texto = ""
    contador = contador_inicial
    for _, row in df.iterrows():
        # Solo mostrar los que NO cumplen
        if row['Cumplimiento'] == "No Cumple":
            # Convertir el contador a número romano
            num_romano = int_to_roman(contador)
            observaciones_texto += (
                f'{num_romano}. {row["Observación"]}\n'
            )
            contador += 1
    return observaciones_texto.strip(), contador


def generar_reporte():
    # Datos principales
    tipo = st.session_state.get("seleccion", "")
    numero_solicitud = st.session_state.get("numero_solicitud", "")
    nombre_titular = st.session_state.get("nombre_titular", "")
    nombre_establecimiento = st.session_state.get("nombre_establecimiento", "")
    tipo_tramite = st.session_state.get("tipo_tramite", "")

    # Ajustar descripción según el tipo
    if tipo == "Unidades de transporte de alimentos":
        descripcion_tipo = (
            f"Tipo de trámite: {tipo_tramite}\n"
            f"Tipo de establecimiento: {tipo}\n"
            f"Nombre de titular: {nombre_titular}\n"
            f"Nombre de establecimiento: {nombre_establecimiento}\n"
        )
    else:
        descripcion_tipo = (
            f"Tipo de trámite: {tipo_tramite}\n"
            f"Tipo de establecimiento: {tipo}\n"
            f"Número de solicitud: {numero_solicitud}\n"
            f"Nombre de titular: {nombre_titular}\n"
            f"Nombre de establecimiento: {nombre_establecimiento}\n"
        )

    # Reunir todas las observaciones de todos los documentos
    observaciones = []
    for doc in botones_sidebar:
        requisitos_filtrados = df_requisitos[df_requisitos["Documento"] == doc]
        observaciones_dict = st.session_state["evaluaciones_documentos"].get(doc, {})
        for idx, row in requisitos_filtrados.iterrows():
            requisito = row["Requisito"]
            cumple = observaciones_dict.get((doc, requisito), {}).get("cumple", "")
            observacion = observaciones_dict.get((doc, requisito), {}).get("observacion", "")
            observaciones.append({
                "Sección": doc,
                "Requisito": requisito,
                "Cumplimiento": cumple,
                "Observación": observacion
            })
    df = pd.DataFrame(observaciones)

    # Selección de plantilla según tipo de establecimiento
    if tipo == "Unidades de transporte de alimentos":
        template_path = "Dictamen_unidades_transporte.docx"
    else:
        template_path = "Dictamen_establecimientos.docx"

    documento = Document(template_path)

    # Reemplazar {tabla} SOLO si es unidades de transporte
    if tipo == "Unidades de transporte de alimentos":
        # Encontrar el párrafo con {tabla}
        tabla_idx = None
        for i, paragraph in enumerate(documento.paragraphs):
            if "{tabla}" in paragraph.text:
                tabla_idx = i
                # Borra el texto del placeholder
                paragraph.text = ""
                break
        
        if tabla_idx is not None:
            # Obtenemos el párrafo actual
            current_paragraph = documento.paragraphs[tabla_idx]
            
            # Crear tabla después del párrafo actual (que ahora está vacío)
            tabla = documento.add_table(rows=1, cols=3)
            
            # Mover la tabla desde el final del documento hasta después del párrafo
            elemento_tabla = tabla._element
            elemento_parrafo = current_paragraph._element
            elemento_parrafo.addnext(elemento_tabla)
            
            # Aplicar bordes y estilo
            try:
                tabla.style = 'Table Grid'
            except KeyError:
                try:
                    tabla.style = 'Cuadrícula'
                except KeyError:
                    pass
            set_table_borders(tabla)
            
            # Llenar contenido con formato negrita en encabezados
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'REFERENCIA'
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto

            hdr_cells[1].text = 'VEHÍCULO'
            hdr_cells[1].paragraphs[0].runs[0].bold = True
            hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto

            hdr_cells[2].text = 'FICHA DE INSPECCIÓN TÉCNICA PRESENTADA'
            hdr_cells[2].paragraphs[0].runs[0].bold = True
            hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto

            # Llenar la tabla con los datos de placas_vehiculos
            for par in st.session_state.get("placas_vehiculos", []):
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(par.get("solicitud", ""))
                row_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
                
                row_cells[1].text = str(par.get("placa", ""))
                row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
                
                row_cells[2].text = str(par.get("tipo_establecimiento", ""))
                row_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centrar texto
            
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    
            # Insertar párrafo con espacios después de la tabla
            # Se inserta un párrafo nuevo con espacios entre la tabla y el siguiente placeholder
            espacio_parrafo = documento.add_paragraph("\n\n")  # Dos saltos de línea
            elemento_espacio = espacio_parrafo._element
            tabla._element.addnext(elemento_espacio)
    
    # Reemplazar {tipo}
    for paragraph in documento.paragraphs:
        if "{tipo}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{tipo}", descripcion_tipo.strip())

    # Reemplazar {Observaciones} con alineación explícita a la izquierda
    observaciones_texto = ""
    contador_global = 1
    for doc_name in botones_sidebar:
        df_doc = df[df['Sección'] == doc_name]
        obs_text, contador_global = generar_observaciones_texto(df_doc, doc_name, contador_global)
        if obs_text:
            observaciones_texto += obs_text + "\n"

    for paragraph in documento.paragraphs:
        if "{Observaciones}" in paragraph.text:
            # Reemplazar el contenido
            paragraph.text = paragraph.text.replace("{Observaciones}", observaciones_texto.strip())
            # Establecer explícitamente la alineación a la izquierda
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Reemplazar {dictamen}
    for paragraph in documento.paragraphs:
        if "{dictamen}" in paragraph.text:
            # Primero limpiamos el placeholder
            paragraph.text = ""
            
            # Para cada documento, agregamos secciones formateadas
            for doc_name in botones_sidebar:
                df_doc = df[df['Sección'] == doc_name]
                
                # Obtenemos secciones únicas
                secciones = df_doc['Sección'].unique()
                for seccion in secciones:
                    df_seccion = df_doc[df_doc['Sección'] == seccion]
                    df_seccion['Requisito'] = df_seccion['Requisito'].apply(limpiar_requisito)
                    
                    # Título en mayúsculas y negrita
                    run = paragraph.add_run(f"{seccion.upper()}\n\n")
                    run.bold = True
                    
                    # Requisitos que cumplen
                    cumple = df_seccion[df_seccion['Cumplimiento'] == "Cumple"]['Requisito'].tolist()
                    if cumple:
                        paragraph.add_run(f"Luego de la evaluación se determina que CUMPLEN los siguientes requisitos: {', '.join(cumple)}.\n\n")
                    
                    # Requisitos que no cumplen
                    no_cumple = df_seccion[df_seccion['Cumplimiento'] == "No Cumple"]['Requisito'].tolist()
                    if no_cumple:
                        paragraph.add_run(f"Los requisitos que NO CUMPLEN son: {', '.join(no_cumple)}.\n\n")
                    
                    # Requisitos no evaluados
                    no_evaluados = df_seccion[df_seccion['Cumplimiento'].isnull() | (df_seccion['Cumplimiento'] == "")]['Requisito'].tolist()
                    if no_evaluados:
                        paragraph.add_run(f"Y los requisitos que NO HAN SIDO EVALUADOS son: {', '.join(no_evaluados)}.\n\n")

    # Reemplazar {TITULAR} en la plantilla de unidades de transporte
    if tipo == "Unidades de transporte de alimentos":
        for paragraph in documento.paragraphs:
            if "{TITULAR}" in paragraph.text:
                nombre_titular = st.session_state.get("nombre_titular", "")
                paragraph.text = paragraph.text.replace("{TITULAR}", nombre_titular)

    # Reemplazar {num_solicitud} y {NOMBRE TRAMITE} para la plantilla de establecimientos
    if tipo != "Unidades de transporte de alimentos":  # Para la plantilla de establecimientos
        # Reemplazar {num_solicitud} con formato negrita
        for paragraph in documento.paragraphs:
            if "{num_solicitud}" in paragraph.text:
                text_before = paragraph.text.split("{num_solicitud}")[0]
                text_after = paragraph.text.split("{num_solicitud}")[1] if len(paragraph.text.split("{num_solicitud}")) > 1 else ""
                
                # Limpiar el párrafo actual
                paragraph.clear()
                
                # Agregar el texto anterior al placeholder
                if text_before:
                    paragraph.add_run(text_before)
                
                # Agregar el número de solicitud en negrita
                numero_solicitud = st.session_state.get("numero_solicitud", "")
                bold_run = paragraph.add_run(numero_solicitud)
                bold_run.bold = True
                
                # Agregar el texto posterior al placeholder
                if text_after:
                    paragraph.add_run(text_after)
        
        # Reemplazar {NOMBRE TRAMITE} con formato negrita
        for paragraph in documento.paragraphs:
            if "{NOMBRE TRAMITE}" in paragraph.text:
                text_before = paragraph.text.split("{NOMBRE TRAMITE}")[0]
                text_after = paragraph.text.split("{NOMBRE TRAMITE}")[1] if len(paragraph.text.split("{NOMBRE TRAMITE}")) > 1 else ""
                
                # Limpiar el párrafo actual
                paragraph.clear()
                
                # Agregar el texto anterior al placeholder
                if text_before:
                    paragraph.add_run(text_before)
                
                # Agregar el tipo de trámite en negrita
                tipo_tramite = st.session_state.get("tipo_tramite", "").upper()
                bold_run = paragraph.add_run(tipo_tramite)
                bold_run.bold = True
                
                # Agregar el texto posterior al placeholder
                if text_after:
                    paragraph.add_run(text_after)

    # Guardar el documento generado
    output = io.BytesIO()
    documento.save(output)
    output.seek(0)
    return output

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)

def generar_dictamen_complementario():
    # Datos principales
    tipo = st.session_state.get("seleccion", "")
    numero_solicitud = st.session_state.get("numero_solicitud", "")
    nombre_titular = st.session_state.get("nombre_titular", "")
    nombre_establecimiento = st.session_state.get("nombre_establecimiento", "")
    tipo_tramite = st.session_state.get("tipo_tramite", "")

    # Ajustar descripción según el tipo
    if tipo == "Unidades de transporte de alimentos":
        descripcion_tipo = (
            f"Tipo de trámite: {tipo_tramite}\n"
            f"Tipo de establecimiento: {tipo}\n"
            f"Nombre de titular: {nombre_titular}\n"
            f"Nombre de establecimiento: {nombre_establecimiento}\n"
        )
    else:
        descripcion_tipo = (
            f"Tipo de trámite: {tipo_tramite}\n"
            f"Tipo de establecimiento: {tipo}\n"
            f"Número de solicitud: {numero_solicitud}\n"
            f"Nombre de titular: {nombre_titular}\n"
            f"Nombre de establecimiento: {nombre_establecimiento}\n"
        )

    # Reunir todas las observaciones de todos los documentos
    observaciones = []
    for doc in botones_sidebar:
        requisitos_filtrados = df_requisitos[df_requisitos["Documento"] == doc]
        observaciones_dict = st.session_state["evaluaciones_documentos"].get(doc, {})
        for idx, row in requisitos_filtrados.iterrows():
            requisito = row["Requisito"]
            cumple = observaciones_dict.get((doc, requisito), {}).get("cumple", "")
            observacion = observaciones_dict.get((doc, requisito), {}).get("observacion", "")
            observaciones.append({
                "Sección": doc,
                "Requisito": requisito,
                "Cumplimiento": cumple,
                "Observación": observacion
            })
    df = pd.DataFrame(observaciones)

    # Usar la plantilla DICTAMEN.docx
    template_path = "DICTAMEN.docx"

    # Cargar el documento
    try:
        documento = Document(template_path)
    except Exception as e:
        st.error(f"Error al cargar la plantilla DICTAMEN.docx: {e}")
        return None

    # Reemplazar {tipo}
    for paragraph in documento.paragraphs:
        if "{tipo}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{tipo}", descripcion_tipo.strip())

    # Reemplazar {dictamen}
    for paragraph in documento.paragraphs:
        if "{dictamen}" in paragraph.text:
            # Primero limpiamos el placeholder
            paragraph.text = ""
            
            # Para cada documento, agregamos secciones formateadas
            for doc_name in botones_sidebar:
                df_doc = df[df['Sección'] == doc_name]
                
                # Obtenemos secciones únicas
                secciones = df_doc['Sección'].unique()
                for seccion in secciones:
                    df_seccion = df_doc[df_doc['Sección'] == seccion]
                    df_seccion['Requisito'] = df_seccion['Requisito'].apply(limpiar_requisito)
                    
                    # Título en mayúsculas y negrita
                    run = paragraph.add_run(f"{seccion.upper()}\n\n")
                    run.bold = True
                    
                    # Requisitos que cumplen
                    cumple = df_seccion[df_seccion['Cumplimiento'] == "Cumple"]['Requisito'].tolist()
                    if cumple:
                        paragraph.add_run(f"Luego de la evaluación se determina que CUMPLEN los siguientes requisitos: {', '.join(cumple)}.\n\n")
                    
                    # Requisitos que no cumplen
                    no_cumple = df_seccion[df_seccion['Cumplimiento'] == "No Cumple"]['Requisito'].tolist()
                    if no_cumple:
                        paragraph.add_run(f"Los requisitos que NO CUMPLEN son: {', '.join(no_cumple)}.\n\n")
                    
                    # Requisitos no evaluados
                    no_evaluados = df_seccion[df_seccion['Cumplimiento'].isnull() | (df_seccion['Cumplimiento'] == "")]['Requisito'].tolist()
                    if no_evaluados:
                        paragraph.add_run(f"Y los requisitos que NO HAN SIDO EVALUADOS son: {', '.join(no_evaluados)}.\n\n")

    # Guardar el documento generado
    output = io.BytesIO()
    documento.save(output)
    output.seek(0)
    return output

# Botón para generar y descargar el informe
if st.sidebar.button('Generar Reportes', key="generar", type='secondary'):
    # Guardar la evaluación actual antes de generar el reporte
    if documento_actual:
        st.session_state["evaluaciones_documentos"][documento_actual] = st.session_state["observaciones_dict"].copy()
    
    # Generar el reporte principal
    output = generar_reporte()
    
    # Generar el dictamen complementario
    output_dictamen = generar_dictamen_complementario()
    
    # Determinar el nombre del archivo para el reporte principal
    tipo = st.session_state.get("seleccion", "")
    if tipo == "Unidades de transporte de alimentos":
        nombre_archivo = f"{st.session_state.get('nombre_establecimiento', '')}.docx" if st.session_state.get("nombre_establecimiento", "") else "Evaluación unidades de transporte.docx"
        nombre_dictamen = f"Dictamen_{st.session_state.get('nombre_establecimiento', '')}.docx" if st.session_state.get("nombre_establecimiento", "") else "Dictamen_unidades_transporte.docx"
    else:
        numero_solicitud = st.session_state.get("numero_solicitud", "")
        nombre_archivo = f"{numero_solicitud}.docx" if numero_solicitud else "Reporte.docx"
        nombre_dictamen = f"Dictamen_{numero_solicitud}.docx" if numero_solicitud else "Dictamen_reporte.docx"
    
    # Crear contenedor para organizar los botones de descarga
    descarga_col1, descarga_col2 = st.sidebar.columns(2)
    
    with descarga_col1:
        st.download_button(
            "Descargar Observaciones",
            output,
            nombre_archivo,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="descargar",
            type='secondary'
        )
    
    with descarga_col2:
        if output_dictamen:  # Asegurarse de que el dictamen se generó correctamente
            st.download_button(
                "Descargar Dictamen",
                output_dictamen,
                nombre_dictamen,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="descargar_dictamen",
                type='secondary'
            )